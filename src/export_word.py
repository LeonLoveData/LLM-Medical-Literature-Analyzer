from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.config import DRUG_CONFIG, OUTPUT_DIR
from src.charts import chart_clinical_phase_distribution, chart_modality_landscape
from src.validation import REQUIRED_FIELDS


def export_word(records, summaries, analysis):
    """Generate BD-ready Word report."""
    doc = Document()

    # Title
    title = doc.add_heading("", level=0)
    run = title.add_run(
        f"{DRUG_CONFIG['drug_name'].replace('_', ' ')} – Clinical Trial & Patent Analysis"
    )
    run.font.color.rgb = RGBColor(0x6A, 0x1D, 0x8A)

    doc.add_paragraph(
        f"Indication: {DRUG_CONFIG['disease']}\n"
        f"Company: {DRUG_CONFIG['company']}\n"
        f"Focus Area: {DRUG_CONFIG['focus_area']}\n"
        f"Data Source: PubMed + Lens Patents\n"
        f"Structured Fields: Extracted with source sentences\n"
        f"Summary & Analysis: LLM inferred\n"
        f"Purpose: BD / Strategy / Due Diligence"
    )

    # Charts
    doc.add_heading("Visual Intelligence Overview", level=1)

    doc.add_heading("Figure 1 – Clinical Phase Distribution", level=2)
    buf1 = chart_clinical_phase_distribution(records)
    doc.add_picture(buf1, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Figure 2 – Therapeutic Modality Scorecard", level=2)
    buf2 = chart_modality_landscape(records)
    doc.add_picture(buf2, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Structured evidence
    doc.add_heading("Structured Evidence", level=1)
    for r in records:
        doc.add_paragraph(f"Source ID: {r['source_id']}", style="Intense Quote")
        for f in REQUIRED_FIELDS:
            v = r.get(f)
            if v and isinstance(v, dict):
                doc.add_paragraph(
                    f"{f.upper()}:\n"
                    f"  Value: {v.get('value')}\n"
                    f"  Source: {v.get('source_text')}"
                )

    # Summaries
    doc.add_heading("LLM Inferred Summaries", level=1)
    for s in summaries:
        doc.add_paragraph(s)

    # Competitive analysis
    doc.add_heading("LLM Inferred Competitive Analysis", level=1)
    doc.add_paragraph(analysis)

    path = f"{OUTPUT_DIR}/{DRUG_CONFIG['drug_name']}_BD_READY_REPORT.docx"
    doc.save(path)
    return path
